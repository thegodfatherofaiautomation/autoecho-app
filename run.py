from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import argparse
import os
import uuid
import subprocess
import whisper

# Generate unique ID
def generate_user_id():
    return str(uuid.uuid4())

# Calculate duration in minutes
def get_audio_duration(file_path):
    command = [
        "ffprobe", "-v", "error", "-show_entries",
        "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", file_path
    ]
    result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    duration = float(result.stdout.strip())
    return duration / 60

# Build transcript document
def generate_transcript_docx(transcript_text, output_path, audio_filename, tier, logo_path=os.path.join("static", "Logo.png")):
    document = Document()

    # Set font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header with logo on top-left
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    if os.path.exists(logo_path):
        run.add_picture(logo_path, width=Inches(1.0))
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Title
    title = document.add_heading(f"Transcription: {audio_filename}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Transcribed text
    para = document.add_paragraph(transcript_text)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Watermark (footer)
    if tier.lower() in ["free", "basic"]:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "AutoEcho Free/Basic Tier – This document was auto-generated."
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save(output_path)
    print(f"[✓] Transcript saved to {output_path}")

# Main function
def transcribe_audio(tier, input_path):
    audio_filename = os.path.basename(input_path)
    user_id = generate_user_id()

    tier_limits = {
        "free": 30,
        "basic": 180,
        "standard": 1200,
        "premium": float("inf"),
        "enterprise": float("inf")
    }

    # Duration check
    duration = get_audio_duration(input_path)
    limit = tier_limits.get(tier.lower(), 0)
    if duration > limit:
        print(f"[!] Exceeds tier limit ({limit} min). Your file is {duration:.2f} min.")
        return

    print(f"[•] Transcribing {audio_filename} ({duration:.2f} min)...")

    # Load Whisper
    model = whisper.load_model("base")
    result = model.transcribe(input_path)
    transcript_text = result["text"]

    # Save .docx
    output_path = f"output/AutoEcho_Transcript_{audio_filename.replace('.mp3','')}.docx"
    generate_transcript_docx(transcript_text, output_path, audio_filename, tier)

# Entry
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--tier", required=True, help="Tier: Free, Basic, Standard, Premium, Enterprise")
    parser.add_argument("--input", required=True, help="Input audio file path")
    args = parser.parse_args()

    transcribe_audio(args.tier, args.input)
