import stripe
from flask import Flask, render_template, request, redirect, url_for, send_file, jsonify
import os
import uuid
import whisper
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB limit

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Stripe keys from environment
stripe.api_key = os.environ.get("STRIPE_SECRET_KEY")
STRIPE_PRICE_BASIC = os.environ.get("STRIPE_PRICE_BASIC")
STRIPE_PRICE_STANDARD = os.environ.get("STRIPE_PRICE_STANDARD")
STRIPE_PRICE_PREMIUM = os.environ.get("STRIPE_PRICE_PREMIUM")
DOMAIN = os.environ.get("DOMAIN", "https://autoecho.xyz")

# Load Whisper model
model = whisper.load_model("tiny")

@app.route("/")
def index():
    return render_template("landing.html")

@app.route("/login")
def login():
    return render_template("login.html")

@app.route("/login/stripe")
def login_stripe():
    try:
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            mode="subscription",
            line_items=[{
                "price": STRIPE_PRICE_BASIC,
                "quantity": 1,
            }],
            success_url=f"{DOMAIN}/upload?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{DOMAIN}/",
        )
        return redirect(checkout_session.url, code=303)
    except Exception as e:
        return jsonify(error=str(e)), 400

@app.route("/upload")
def upload():
    return render_template("upload.html")

@app.route("/transcribe", methods=["POST"])
def transcribe():
    if "file" not in request.files:
        return "No file uploaded", 400

    file = request.files["file"]
    if file.filename == "":
        return "Empty filename", 400

    tier = request.form.get("tier", "Free").capitalize()
    filename = f"{uuid.uuid4().hex}_{file.filename}"
    input_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(input_path)

    try:
        result = model.transcribe(input_path)
        transcript_text = result["text"]
    except Exception as e:
        return f"Transcription failed: {str(e)}", 500

    doc_filename = f"Transcript_{os.path.splitext(file.filename)[0]}.docx"
    doc_path = os.path.join(OUTPUT_FOLDER, doc_filename)
    generate_transcript_docx(transcript_text, doc_path, file.filename, tier)

    return send_file(doc_path, as_attachment=True)

@app.errorhandler(413)
def request_entity_too_large(error):
    return "File too large. Max upload size is 20MB.", 413

def generate_transcript_docx(transcript_text, output_path, audio_filename, tier, logo_path=os.path.join("static", "Logo.png")):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    if os.path.exists(logo_path):
        run.add_picture(logo_path, width=Inches(1.0))
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    title = document.add_heading(f"Transcription: {audio_filename}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    para = document.add_paragraph(transcript_text)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    if tier.lower() in ["free", "basic"]:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "AutoEcho Free/Basic Tier – This document was auto-generated. English transcriptions only."
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save(output_path)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)
